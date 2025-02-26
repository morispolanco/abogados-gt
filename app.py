import streamlit as st
import pandas as pd
from datetime import datetime
import sqlite3
import requests
import json
from docx import Document

# Configuración inicial para modo ancho
st.set_page_config(page_title="Gestor Legal GT")

# --- Configuración de la Base de Datos SQLite ---
DB_FILE = "gestor_legal.db"

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS casos 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, cliente TEXT, tipo TEXT, 
                  fecha_inicio TEXT, estado TEXT)''')
    conn.commit()
    conn.close()

init_db()

# --- Configuración de la API Key en secrets ---
API_KEY = st.secrets.get("api_key", None)
if not API_KEY:
    st.error("API Key no configurada. Agrega 'api_key' en .streamlit/secrets.toml")
    st.stop()

# --- Función para llamar a la API de Google Gemini ---
def generate_legal_content(prompt):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={API_KEY}"
    headers = {"Content-Type": "application/json"}
    payload = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 1,
            "topK": 40,
            "topP": 0.95,
            "maxOutputTokens": 8192,
            "responseMimeType": "text/plain"
        }
    }
    response = requests.post(url, headers=headers, data=json.dumps(payload))
    if response.status_code == 200:
        content = response.json()["candidates"][0]["content"]["parts"][0]["text"]
        content = content.replace("**", "").replace("*", "").replace("#", "").replace("`", "")
        return content
    else:
        st.error(f"Error en la API: {response.text}")
        return "Contenido no generado debido a un error."

# --- Barra Lateral con Explicación ---
st.sidebar.header("¿Qué es Gestor Legal GT?")
st.sidebar.markdown("""
**Gestor Legal GT** es una herramienta diseñada para abogados en Guatemala. Simplifica tu trabajo diario al permitirte:

- **Gestionar casos**: Registra y organiza tus casos con facilidad.
- **Calcular honorarios**: Obtén tarifas precisas en segundos, con o sin IVA.
- **Generar documentos**: Crea demandas, contratos y recibos ajustados a las leyes guatemaltecas, listos para descargar en Word.

Todo por una licencia única de $500. ¡Optimiza tu práctica legal hoy!
""")

# --- Interfaz Principal ---
st.title("Gestor Legal GT")
st.markdown("**Herramienta para abogados en Guatemala**")

# --- Tabs ---
tab1, tab2, tab3 = st.tabs(["Gestión de Casos", "Cálculo de Honorarios", "Documentos"])

with tab1:
    st.header("Gestión de Casos")
    col1, col2 = st.columns(2)
    with col1:
        with st.form(key="nuevo_caso"):
            cliente = st.text_input("Nombre del Cliente")
            tipo_caso = st.selectbox("Tipo de Caso", ["Civil", "Penal", "Laboral", "Mercantil"])
            fecha_inicio = st.date_input("Fecha de Inicio", datetime.today())
            estado = st.selectbox("Estado", ["En Progreso", "Ganado", "Perdido"])
            submit_button = st.form_submit_button(label="Agregar Caso")
            if submit_button:
                conn = sqlite3.connect(DB_FILE)
                c = conn.cursor()
                c.execute("INSERT INTO casos (cliente, tipo, fecha_inicio, estado) VALUES (?, ?, ?, ?)",
                          (cliente, tipo_caso, str(fecha_inicio), estado))
                conn.commit()
                conn.close()
                st.success("Caso agregado con éxito")
    with col2:
        st.subheader("Lista de Casos")
        conn = sqlite3.connect(DB_FILE)
        casos_df = pd.read_sql_query("SELECT * FROM casos", conn)
        conn.close()
        st.dataframe(casos_df)

with tab2:
    st.header("Cálculo de Honorarios")
    col1, col2 = st.columns(2)
    with col1:
        horas = st.number_input("Horas trabajadas", min_value=1, value=10)
        tarifa_hora = st.number_input("Tarifa por hora (Q)", min_value=50.0, value=150.0, step=10.0)
    with col2:
        incluir_iva = st.checkbox("Incluir IVA (12%)")
        subtotal = horas * tarifa_hora
        iva = subtotal * 0.12 if incluir_iva else 0
        total = subtotal + iva
        st.write(f"Subtotal: Q{subtotal:.2f}")
        if incluir_iva:
            st.write(f"IVA (12%): Q{iva:.2f}")
        st.write(f"**Total: Q{total:.2f}**")

with tab3:
    st.header("Generar Documentos")
    doc_types = [
        "Recibo de Honorarios", "Contrato Privado", "Contrato de Arrendamiento", "Contrato de Compraventa",
        "Contrato de Prestación de Servicios", "Contrato de Mutuo", "Contrato de Sociedad", "Poder General",
        "Escrito de Amparo", "Contrato de Trabajo", "Contrato de Donación", "Contrato de Hipoteca",
        "Demanda Inicial", "Demanda de Desalojo", "Demanda de Divorcio", "Demanda Laboral por Despido Injustificado",
        "Demanda Penal por Estafa", "Demanda de Alimentos", "Demanda de Reconocimiento de Unión de Hecho",
        "Demanda de Nulidad de Contrato", "Demanda de Pago por Cheque sin Fondos", "Demanda de Daños y Perjuicios",
        "Demanda de Cumplimiento de Contrato", "Demanda de Usucapión"
    ]
    doc_type = st.selectbox("Tipo de Documento", doc_types)

    if doc_type == "Recibo de Honorarios":
        nombre_cliente = st.text_input("Nombre del Cliente")
        monto = st.number_input("Monto (Q)", min_value=0.0)
        if st.button("Generar Recibo"):
            doc = Document()
            doc.add_heading("Recibo de Honorarios", 0)
            doc.add_paragraph(f"Cliente: {nombre_cliente}")
            doc.add_paragraph(f"Monto: Q{monto:.2f}")
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
            file_name = f"recibo_{nombre_cliente}_{datetime.now().strftime('%Y%m%d')}.docx"
            doc.save(file_name)
            st.success(f"Recibo generado: {file_name}")
            with open(file_name, "rb") as file:
                st.download_button("Descargar", file, file_name=file_name)

    else:
        parte1 = st.text_input("Nombre de la Parte 1 (DPI si aplica)")
        parte2 = st.text_input("Nombre de la Parte 2 (DPI si aplica)")
        if doc_type.startswith("Contrato"):
            objeto = st.text_area("Objeto del Contrato")
            monto_contrato = st.number_input("Monto del Contrato (Q)", min_value=0.0)
        elif doc_type == "Poder General":
            objeto = "Otorgamiento de poder general a favor de la Parte 2"
            monto_contrato = 0.0
        elif doc_type == "Escrito de Amparo":
            motivo = st.text_area("Motivo del Amparo")
            pretension = st.text_input("Pretensión")
        elif doc_type.startswith("Demanda"):
            motivo = st.text_area("Motivo de la Demanda")
            pretension = st.text_input("Pretensión")

        if st.button(f"Generar {doc_type}"):
            doc = Document()
            doc.add_heading(doc_type, 0)
            doc.add_paragraph(f"Parte 1: {parte1}")
            doc.add_paragraph(f"Parte 2: {parte2}")

            if doc_type == "Escrito de Amparo":
                prompt = f"Redacta un escrito de amparo conforme a la Ley de Amparo de Guatemala. Recurrente: {parte1}, Autoridad: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye estructura formal."
            elif doc_type == "Demanda Inicial":
                prompt = f"Redacta una demanda inicial genérica conforme al Código Procesal Civil y Mercantil de Guatemala. Demandante: {parte1}, Demandado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye estructura formal y referencia a leyes guatemaltecas."
            elif doc_type == "Demanda de Desalojo":
                prompt = f"Redacta una demanda de desalojo conforme al Código Civil de Guatemala. Demandante (arrendador): {parte1}, Demandado (arrendatario): {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye referencia a leyes aplicables."
            elif doc_type == "Demanda de Divorcio":
                prompt = f"Redacta una demanda de divorcio conforme al Código Civil de Guatemala. Demandante: {parte1}, Demandado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye estructura formal."
            elif doc_type == "Demanda Laboral por Despido Injustificado":
                prompt = f"Redacta una demanda laboral por despido injustificado conforme al Código de Trabajo de Guatemala. Demandante (trabajador): {parte1}, Demandado (empleador): {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye referencias legales."
            elif doc_type == "Demanda Penal por Estafa":
                prompt = f"Redacta una demanda penal por estafa conforme al Código Penal de Guatemala. Querellante: {parte1}, Querellado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye estructura formal."
            elif doc_type == "Demanda de Alimentos":
                prompt = f"Redacta una demanda de alimentos conforme al Código Civil de Guatemala. Demandante: {parte1}, Demandado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye referencias legales."
            elif doc_type == "Demanda de Reconocimiento de Unión de Hecho":
                prompt = f"Redacta una demanda de reconocimiento de unión de hecho conforme al Código Civil de Guatemala. Demandante: {parte1}, Demandado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye estructura formal."
            elif doc_type == "Demanda de Nulidad de Contrato":
                prompt = f"Redacta una demanda de nulidad de contrato conforme al Código Civil de Guatemala. Demandante: {parte1}, Demandado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye referencias legales."
            elif doc_type == "Demanda de Pago por Cheque sin Fondos":
                prompt = f"Redacta una demanda de pago por cheque sin fondos conforme al Código de Comercio de Guatemala. Demandante: {parte1}, Demandado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye estructura formal."
            elif doc_type == "Demanda de Daños y Perjuicios":
                prompt = f"Redacta una demanda de daños y perjuicios conforme al Código Civil de Guatemala. Demandante: {parte1}, Demandado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye referencias legales."
            elif doc_type == "Demanda de Cumplimiento de Contrato":
                prompt = f"Redacta una demanda de cumplimiento de contrato conforme al Código Civil de Guatemala. Demandante: {parte1}, Demandado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye estructura formal."
            elif doc_type == "Demanda de Usucapión":
                prompt = f"Redacta una demanda de usucapión (prescripción adquisitiva) conforme al Código Civil de Guatemala. Demandante: {parte1}, Demandado: {parte2}, Motivo: '{motivo}', Pretensión: '{pretension}'. Incluye referencias legales."
            else:
                prompt = f"Redacta un {doc_type} conforme a las leyes de Guatemala entre {parte1} y {parte2}, con el objeto: '{objeto}', por un monto de Q{monto_contrato}. Incluye cláusulas estándar como cumplimiento, resolución y jurisdicción en Guatemala."

            contenido = generate_legal_content(prompt)
            doc.add_paragraph(contenido)

            file_name = f"{doc_type.lower().replace(' ', '_')}_{parte1}_{datetime.now().strftime('%Y%m%d')}.docx"
            doc.save(file_name)
            st.success(f"Documento generado: {file_name}")
            with open(file_name, "rb") as file:
                st.download_button("Descargar", file, file_name=file_name)

# Instrucciones
st.sidebar.markdown("**Ejecutar:** `streamlit run app.py`")
